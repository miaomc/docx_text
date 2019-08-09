# -*- coding:cp936 -*-
import docx
import win32com.client
import os
import json
    

def init(docName):
    """
    ����docx�Ķ���
    """
    return docx.Document(docName)
    
def read_paragraphs(doc):
    """��ȡdocx�����еĶ���"""
    fullText = []
    paras = doc.paragraphs
    for p in paras:
        fullText.append(p.text)
    return '\n'.join(fullText)

def read_tables(doc):
    """��ȡdocx�����еı��"""
    tblText = []
    for t in doc.tables:
        for row in t.rows:
            tmp = ''
            for cell in row.cells:
                tmp += read_paragraphs(cell) + '\t'
                tblText.append(tmp)

    return '\n'.join(tblText)


def replace_text(doc, old_text, new_text):
    """�滻docx�����е��ı����ݣ�����Դ��ʽ"""
    for p in doc.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in inline:
                if old_text in i.text:
                    text = i.text.replace(old_text, new_text)
                    i.text = text

def doc2docx(old_file, new_file):
    """��*.doc�ļ�ת��Ϊ*.docx�ļ���Ŀǰ��֪*.doc�ļ�����ֻ�е�ȡ�Դ���com�ӿ�"""
    app=win32com.client.Dispatch('Word.Application')
    doc=app.Documents.Open(old_file)
    doc.SaveAs(new_file, 16)
    doc.Close()
    app.Quit()

def search_file(search_path):
    """���ұ���������Ŀ¼�µ������ļ�"""
    lst_caiye = []
    for path, dirs, files in os.walk(search_path):
        for f in files:
             print f
             lst_caiye.append(os.path.join(path, f))

    return lst_caiye


# ----------------�ļ���������-------------------------------

def test_word():
    """���ԣ���ȡdocx�ļ��е�����"""
    docName = u'��ҳV1.1.docx'
    doc = init(docName)
    print read_paragraphs(doc)
    print read_tables(doc)

def make_doxc2json(path):
    """������*.doc���ļ�ת��Ϊ*.docx, ��ȡ���ݣ�����{�ļ������ļ�����}��json """
    # make all *.doc SaveAs *.docx
    path = raw_input('Default ColorPage Path('+path+'):') or path
    l = search_file(path)
    for i in l:
        if i.endswith('.doc'):
            if i+'x' not in l: # if it has not been done before
                print i+' been to .docx...'
                doc2docx(i,i+'x')

    # make all to MEMORY
    import time
    t1 =  time.time()
    d = {}
    l = search_file(path)
    for i in l:
        if i.endswith('.docx'):
            try:
                doc = init(i)
                print i
                d[i] = read_paragraphs(doc) + read_tables(doc)
            except:
                print i,' open failed...'
    print time.time() - t1

    with open('products.json','w') as f1:
        json.dump(d,fp=f1,encoding='cp936')


# ---------------����json�еĹؼ��֣������ļ���----------------

def search_text(d, key_value):
    """�����ɵ�json���Ǹ��ֵ�{�ļ������ļ�����}���в��Ҷ�Ӧ�ļ��Ĺؼ��֣�
    �����ļ����б�"""
    lst_res = []
    for i in d.keys():
        if key_value in d[i]:
            start = d[i].index(key_value)-10
            if start<0:
                start = 0
            tmp = [i,d[i][start:start+20]]
            lst_res.append(tmp)
    return lst_res
        
def search_in_json():
    """
    �����ɵ�json���Ǹ��ֵ�{�ļ������ļ�����}���в��Ҷ�Ӧ�ļ��Ĺؼ���;
    �����ÿո�����������ؼ���
    """
    with open('products.json','r') as f1:
        d = json.load(f1,encoding='cp936')
    while True:
        text = raw_input("Enter 'q' to quit:").decode('cp936')
        if text == 'q':
            return

        text_list = text.split(' ')  # �ÿո����
        result ={}  # {key1:[[filename11, content11], [filename12, content12]...]...}
        for key_tmp in text_list:
            result[key_tmp] = search_text(d, key_value = key_tmp)

##        return result

        this_key = result.keys()[0]
        for filename_content in result[this_key]:
            filename = filename_content[0]
            whether_print = True
            for other_key in result.keys():
                if this_key==other_key:
                    continue
                if  filename not in [f_c[0] for f_c in result[other_key]]:
                    whether_print = False
                    break
            if whether_print:
                print filename
                for key_tmp in result.keys():
                    for f_c_tmp in result[key_tmp]:
                        if f_c_tmp[0]==filename:
                            print '[',f_c_tmp[1],']'
        print            
        
if __name__ == '__main__':
##    make_doxc2json(path = r'D:\python\docx_text\colorpage_finder\201908����IPC&NVR��ҳ')
    tmp = search_in_json()
    
##  ����һ�� "~$V IPC322E-H 1080P����������������������ҳV2.1.docx"�ļ�
##    import glob
##    l = glob.glob(r'F:\������Ʒ��ҳ\������Ʒ��ҳ(update_20180815)\����\*.docx')
##    print len(l)
##    for i in l:
##        print i
