# -*- coding:cp936 -*-
import docx

def readDocx(docName):
    fullText = []
    doc = docx.Document(docName)
    paras = doc.paragraphs
    for p in paras:
        fullText.append(p.text)
    return '\n'.join(fullText)

def replace_text(doc, old_text, new_text):
    for p in doc.paragraphs:
        if old_text in p.text:
            print 'True in it'
            inline = p.runs
            for i in inline:
                if old_text in i.text:
                    print 'CHANGE:',old_text,new_text
                    text = i.text.replace(old_text, new_text)
                    i.text = text
                    
def read_tables(doc):
    """��ȡdocx�����еı��
        doc.tables[0].rows[0].cells[0].paragraphs[0].runs"""
    tblText = []
    for t in doc.tables:
        print t
        for row in t.rows:
            tmp = ''
            print '\t',row
            for cell in row.cells:
                print '\t\t',cell,cell.text
                tmp += cell.text + '\t'
                tblText.append(tmp)

    return '\n'.join(tblText)

if __name__ == '__main__':
##    docName = u'���쳧����Ȩ����Uniview����ʱЧ������20160101_template.docx'
##    doc = docx.Document(docName)
##    for i,j in enumerate(doc.paragraphs):
##        print i,j.text
##    t = read_tables(doc)

##    replace_text(doc,u'JIAFAGN_NAME',u'XXX����')
##    doc.save('2321.docx')
    import sys
    key = raw_input("Please input a key: ").decode(sys.stdin.encoding)
    print sys.stdin.encoding
