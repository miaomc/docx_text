# -*- coding:cp936 -*-
import docx


def replace_text(doc, old_text, new_text):
    for p in doc.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in inline:
                if old_text in i.text:
                    text = i.text.replace(old_text, new_text)
                    i.text = text
                    
def replace_table(doc, old_text, new_text):
    """�滻docx�����еı��"""
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                replace_text(cell, old_text, new_text)
                
def replace(old_name, new_name, ziku=None):
    doc = docx.Document(old_name)
    for key in ziku.keys():
        replace_text(doc, key, ziku[key])
        replace_table(doc, key, ziku[key])
    doc.save(new_name)
    
def shouquan_fuwu_replace(replace_dir):
    shouquan_docx = u'���쳧����Ȩ����Uniview����ʱЧ������20160101_template.docx'
    shouhou_docx = u'�ۺ�����ŵģ�塾Uniview��20160901_template.docx'

    docx_list = [shouquan_docx, shouhou_docx]
    for i in docx_list:
        if u'���쳧����Ȩ��' in i:
            new_doc_name = u'���쳧����Ȩ��' + '_' + replace_dir['JICHENGSHANG_NAME'] + '.docx'
        elif u'�ۺ�����ŵģ��' in i:
            new_doc_name = u'�ۺ�����ŵģ��' + '_' + replace_dir['JICHENGSHANG_NAME'] + '.docx'
        else:
            new_doc_name = i.split('.')[0] + '_' + replace_dir['JICHENGSHANG_NAME'] + '.docx'
        replace(old_name=i, new_name=new_doc_name, ziku=replace_dir)
