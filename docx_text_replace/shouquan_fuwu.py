# -*- coding:cp936 -*-
import docx
import sys

reload(sys)
sys.setdefaultencoding( "cp936" )

##def readDocx(docName):
##    fullText = []
##    doc = docx.Document(docName)
##    paras = doc.paragraphs
##    for p in paras:
##        fullText.append(p.text)
##    return '\n'.join(fullText)

##def read_tables(doc):
##    """��ȡdocx�����еı��
##        doc.tables[0].rows[0].cells[0].paragraphs[0].runs"""
##    tblText = []
##    for t in doc.tables:
##        print t
##        for row in t.rows:
##            tmp = ''
##            print '\t',row
##            for cell in row.cells:
##                print '\t\t',cell,cell.text
##                tmp += cell.text + '\t'
##                tblText.append(tmp)
##
##    return '\n'.join(tblText)

def replace_text(doc, old_text, new_text):
    for p in doc.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in inline:
                if old_text in i.text:
                    text = i.text.replace(old_text, new_text)
                    i.text = text
                    
def replace_table(doc, old_text, new_text):
    """�滻docx�����еı��"""
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                replace_text(cell, old_text, new_text)
                
def replace(old_name, new_name, ziku=None):
    doc = docx.Document(old_name)
    for key in ziku.keys():
        replace_text(doc, key, ziku[key])
        replace_table(doc, key, ziku[key])
    doc.save(new_name)
    
def shouquan_fuwu_replace(replace_dir):
    shouquan_docx = u'���쳧����Ȩ����Uniview����ʱЧ������20160101_template.docx'
    shouhou_docx = u'��20190709�����ӿƼ��ۺ�����ŵģ��_template.docx'

    docx_list = [shouquan_docx, shouhou_docx]
    for i in docx_list:
        if u'���쳧����Ȩ��' in i:
            new_doc_name = u'���쳧����Ȩ��' + '_' + replace_dir['JICHENGSHANG_NAME'] + '.docx'
        elif u'�ۺ�����ŵģ��' in i:
            new_doc_name = u'�ۺ�����ŵģ��' + '_' + replace_dir['JICHENGSHANG_NAME'] + '.docx'
        else:
            new_doc_name = i.split('.')[0] + '_' + replace_dir['JICHENGSHANG_NAME'] + '.docx'
        replace(old_name=i, new_name=new_doc_name, ziku=replace_dir)

def info_input():
    jiafang_name = u'�׷���XX��λ��'
    xiangmu_name = u'��Ŀ���ƣ�����ʡXX��Ŀ���ɹ���XX����'
    jichengshang_name = u'�����̣�����XX�Ƽ���˾��'
    jichengshang_dizhi = u'�����̵�ַ��ܽ����XX¥XX��XX����XX����'
    time_weibao = u'ά��ʱ�䣨2��'
    time_yyyy = u'�꣨2018��'
    time_mm = u'�£�7��'
    time_dd = u'�գ�22��'

    replace_dir = {u'JIAFANG':jiafang_name,
                   u'XIANGMU':xiangmu_name,
                   u'JICHENGSHANG_NAME':jichengshang_name,
                   u'JICHENGSHANG_DIZHI':jichengshang_dizhi,
                   u'WEIBAO':time_weibao,
                   u'TIME_YYYY':time_yyyy,
                   u'TIME_MM':time_mm,
                   u'TIME_DD':time_dd}
    lst_mannaul = [u'JIAFANG',
                   u'XIANGMU',
                   u'JICHENGSHANG_NAME',
                   u'JICHENGSHANG_DIZHI',
                   u'WEIBAO',
                   u'TIME_YYYY',
                   u'TIME_MM',
                   u'TIME_DD']
    for i in lst_mannaul:
        replace_dir[i] = raw_input(i+'('+replace_dir[i]+'):').decode(sys.stdin.encoding)

    return replace_dir
    
def main():
    d = info_input()
    shouquan_fuwu_replace(d)
    
if __name__ == '__main__':
    main()
