# -*- coding:cp936 -*-
import docx

def readDocx(docName):
    fullText = []
    doc = docx.Document(docName)
    paras = doc.paragraphs
    for p in paras:
        fullText.append(p.text)
    return '\n'.join(fullText)

def replace_text(doc, old_text, new_text):
    for p in doc.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in inline:
                if old_text in i.text:
                    text = i.text.replace(old_text, new_text)
                    i.text = text


if __name__ == '__main__':
    docName = u'11.docx'
    doc = docx.Document(docName)
    replace_text(doc,'A_NAME',u' 哈比BB公司 ')
    doc.save('11.docx')
