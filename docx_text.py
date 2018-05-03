# -*- coding:cp936 -*-
import docx
import win32com.client
import os
import json
    

def init(docName):
    """
    ����docx�Ķ���
    """
    return docx.Document(docName)
    
def read_paragraphs(doc):
    """��ȡdocx�����еĶ���"""
    fullText = []
    paras = doc.paragraphs
    for p in paras:
        fullText.append(p.text)
    return '\n'.join(fullText)

def read_tables(doc):
    """��ȡdocx�����еı���"""
    tblText = []
    for t in doc.tables:
        for row in t.rows:
            tmp = ''
            for cell in row.cells:
                tmp += cell.text + '\t'
        tblText.append(tmp)

    return '\n'.join(tblText)


def replace_text(doc, old_text, new_text):
    """�滻docx�����е��ı����ݣ�����Դ��ʽ"""
    for p in doc.paragraphs:
        if old_text in p.text:
            inline = p.runs
            for i in inline:
                if old_text in i.text:
                    text = i.text.replace(old_text, new_text)
                    i.text = text

def doc2docx(old_file, new_file):
    """��*.doc�ļ�ת��Ϊ*.docx�ļ���Ŀǰ��֪*.doc�ļ�����ֻ�е�ȡ�Դ���com�ӿ�"""
    app=win32com.client.Dispatch('Word.Application')
    doc=app.Documents.Open(old_file)
    doc.SaveAs(new_file, 16)
    doc.Close()
    app.Quit()

def search_file(search_path = r'F:\������Ʒ��ҳ'):
    """���ұ���������Ŀ¼�µ������ļ�"""
    lst_caiye = []
    for path, dirs, files in os.walk(search_path):
        for f in files:
             lst_caiye.append(os.path.join(path, f))

    return lst_caiye


# ----------------�ļ�����������-------------------------------

def test_word():
    """���ԣ���ȡdocx�ļ��е�����"""
    docName = u'��ҳV1.1.docx'
    doc = init(docName)
    print read_paragraphs(doc)
    print read_tables(doc)

def make_doxc2json():
    """������*.doc���ļ�ת��Ϊ*.docx, ��ȡ���ݣ�����{�ļ������ļ�����}��json """
    # make all *.doc SaveAs *.docx 
    l = search_file()
    for i in l:
        if i.endswith('.doc'):
            if i+'x' not in l: # if it has not been done before
                print i
                doc2docx(i,i+'x')

    # make all to MEMORY
    import time
    t1 =  time.time()
    d = {}
    l = search_file()
    for i in l:
        if i.endswith('.docx'):
            doc = init(i)
            d[i] = read_paragraphs(doc) + read_tables(doc)
    print time.time() - t1

    with open('products.json','w') as f1:
        json.dump(d,fp=f1,encoding='cp936')


# ---------------����json�еĹؼ��֣������ļ���----------------

def search_text(d, key_value):
    """�����ɵ�json���Ǹ��ֵ�{�ļ������ļ�����}���в��Ҷ�Ӧ�ļ��Ĺؼ��֣�
    �����ļ����б�"""
    lst_res = []
    for i in d.keys():
        if key_value in d[i]:
            lst_res.append(i)
    return lst_res
        
def search_in_json():
    
    with open('products.json','r') as f1:
        d = json.load(f1,encoding='cp936')
    while True:
        text = raw_input("Enter 'q' to quit:").decode('cp936')
        if text == 'q':
            exit(0)
        l = search_text(d, key_value = text)
        for i in l:
            print i
        
if __name__ == '__main__':
    # make_doxc2json()
    search_in_json()
